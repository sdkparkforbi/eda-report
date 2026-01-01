"""
EDA Report Generator v2.0
=========================
Complete publication-ready data analysis with Word report generation.

Features:
- Dataset size (rows/columns)
- Variable list with GPT-generated descriptions
- Automatic variable classification (ID/Time/Categorical/Numeric)
- Categorical: frequency table (absolute/relative), pie chart, bar chart
- Numeric: mean, std, percentiles, histogram, boxplot
- Missing value analysis for all variables
- Unique key analysis
- ID/Time variables: format only (excluded from statistics)
- Integrated narrative report with tables and figures

Usage:
    from eda_report import generate_report
    generate_report(df, download=True)
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import zipfile
import warnings
warnings.filterwarnings('ignore')

__version__ = '2.0.0'

# Configuration
DEFAULT_CONFIG = {
    'fig_width': 5,
    'fig_height': 4,
    'fig_dpi': 300,
    'n_representative': 3,
    'colors': {
        'primary': '#2E86AB',
        'secondary': '#A23B72',
        'accent': '#F18F01',
        'light': '#E8E8E8'
    }
}


def _get_openai_client(api_key=None):
    """Get OpenAI client with API key from various sources."""
    from openai import OpenAI
    
    if api_key:
        return OpenAI(api_key=api_key)
    
    if os.environ.get('OPENAI_API_KEY'):
        return OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
    
    try:
        from google.colab import userdata
        return OpenAI(api_key=userdata.get('OPENAI_API_KEY'))
    except:
        pass
    
    raise ValueError("OpenAI API key not found.")


def get_variable_descriptions(df, client):
    """Get GPT-generated descriptions for each variable."""
    var_info = {col: {
        'dtype': str(df[col].dtype),
        'sample': [str(x) for x in df[col].dropna().head(3).tolist()]
    } for col in df.columns}
    
    prompt = f"""Describe each variable in this dataset in one sentence.
Variables: {json.dumps(var_info, indent=2)}

Return JSON only: {{"variable_name": "one sentence description"}}
Be specific and technical."""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=800
    )
    
    try:
        return json.loads(response.choices[0].message.content.strip()
                         .replace("```json", "").replace("```", ""))
    except:
        return {col: "No description available" for col in df.columns}


def classify_variables(df, client):
    """Classify variables using GPT API."""
    var_info = {col: {
        'dtype': str(df[col].dtype),
        'nunique': int(df[col].nunique()),
        'sample': [str(x) for x in df[col].dropna().head(5).tolist()]
    } for col in df.columns}
    
    prompt = f"""Classify each variable into exactly ONE category:
{json.dumps(var_info, indent=2)}

Categories:
- "id": Identifiers for grouping (patient_id, user_id, subject)
- "time": Time-related (date, timestamp, time_idx, day, month, year)
- "categorical": Nominal categories, no meaningful magnitude (gender, blood_type, day_of_week)
- "numeric_discrete": Integers with meaningful magnitude (dosage, count, age, score)
- "numeric_continuous": Continuous measurements (temperature, glucose, weight)

IMPORTANT RULES:
- Variables named 'patient', 'subject', 'id', 'user' = id
- Integer dosage/medication amounts = numeric_discrete (NOT categorical)
- Day of week (0-6) = categorical (no magnitude meaning)

Return JSON only: {{"variable_name": "category"}}"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500
    )
    
    result = json.loads(response.choices[0].message.content.strip()
                       .replace("```json", "").replace("```", ""))
    
    classification = {
        'id_vars': [], 'time_vars': [], 'categorical': [],
        'numeric_discrete': [], 'numeric_continuous': []
    }
    
    mapping = {
        'id': 'id_vars', 'time': 'time_vars', 'categorical': 'categorical',
        'numeric_discrete': 'numeric_discrete', 'numeric_continuous': 'numeric_continuous'
    }
    
    for var, cat in result.items():
        if var in df.columns:
            classification[mapping.get(cat, 'categorical')].append(var)
    
    return classification


def analyze_missing(df):
    """Analyze missing values for all variables."""
    missing_data = []
    for col in df.columns:
        n_missing = df[col].isnull().sum()
        pct_missing = n_missing / len(df) * 100
        missing_data.append({
            'variable': col,
            'n_missing': n_missing,
            'pct_missing': pct_missing,
            'n_valid': len(df) - n_missing
        })
    return pd.DataFrame(missing_data)


def analyze_unique_keys(df, var_class):
    """Analyze potential unique keys."""
    results = []
    
    # Single column keys
    for col in df.columns:
        n_unique = df[col].nunique()
        is_unique = n_unique == len(df)
        results.append({
            'columns': col,
            'n_unique': n_unique,
            'is_unique_key': is_unique,
            'coverage': n_unique / len(df) * 100
        })
    
    # Composite keys (ID + Time)
    if var_class['id_vars'] and var_class['time_vars']:
        for id_col in var_class['id_vars']:
            for time_col in var_class['time_vars']:
                combo = f"{id_col} + {time_col}"
                n_unique = df.groupby([id_col, time_col]).ngroups
                is_unique = n_unique == len(df)
                results.append({
                    'columns': combo,
                    'n_unique': n_unique,
                    'is_unique_key': is_unique,
                    'coverage': n_unique / len(df) * 100
                })
    
    return pd.DataFrame(results)


def compute_categorical_stats(series):
    """Compute statistics for categorical variable."""
    freq = series.value_counts()
    rel_freq = series.value_counts(normalize=True) * 100
    
    stats_df = pd.DataFrame({
        'Category': freq.index,
        'Count': freq.values,
        'Percentage': rel_freq.values.round(2)
    })
    
    return stats_df


def compute_numeric_stats(series):
    """Compute statistics for numeric variable."""
    data = series.dropna()
    
    stats = {
        'N': len(data),
        'Missing': series.isnull().sum(),
        'Missing %': series.isnull().sum() / len(series) * 100,
        'Mean': data.mean(),
        'Std': data.std(),
        'Min': data.min(),
        '25%': data.quantile(0.25),
        '50% (Median)': data.quantile(0.50),
        '75%': data.quantile(0.75),
        'Max': data.max(),
        'Range': data.max() - data.min(),
        'IQR': data.quantile(0.75) - data.quantile(0.25)
    }
    
    return stats


def generate_figures(df, var_class, config, output_dir='.'):
    """Generate all figures."""
    figure_files = []
    colors = config['colors']
    fig_num = 1
    
    # 1. Categorical variables - Pie chart + Bar chart
    for col in var_class['categorical']:
        freq = df[col].value_counts()
        
        # Pie Chart
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        colors_pie = plt.cm.Set3(np.linspace(0, 1, len(freq)))
        wedges, texts, autotexts = ax.pie(freq.values, labels=freq.index.astype(str), 
                                           autopct='%1.1f%%', colors=colors_pie, startangle=90)
        ax.set_title(f'Distribution of {col}')
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_{col}_pie.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figure_files.append({'path': fname, 'type': 'pie', 'variable': col, 'num': fig_num})
        plt.close()
        fig_num += 1
        
        # Bar Chart (Frequency Distribution)
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        bars = ax.bar(freq.index.astype(str), freq.values, color=colors['primary'], edgecolor='black', width=0.6)
        ax.set_xlabel(col)
        ax.set_ylabel('Frequency')
        ax.set_title(f'Frequency Distribution of {col}')
        for bar, val in zip(bars, freq.values):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                   str(val), ha='center', va='bottom', fontsize=9)
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_{col}_bar.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figure_files.append({'path': fname, 'type': 'bar', 'variable': col, 'num': fig_num})
        plt.close()
        fig_num += 1
    
    # 2. Numeric variables - Histogram + Boxplot
    all_numeric = var_class['numeric_discrete'] + var_class['numeric_continuous']
    
    for col in all_numeric:
        data = df[col].dropna()
        
        # Histogram
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        ax.hist(data, bins=20, color=colors['primary'], edgecolor='black', alpha=0.7)
        ax.axvline(data.mean(), color=colors['secondary'], linestyle='--', lw=2, 
                  label=f'Mean: {data.mean():.2f}')
        ax.axvline(data.median(), color=colors['accent'], linestyle='-.', lw=2, 
                  label=f'Median: {data.median():.2f}')
        ax.set_xlabel(col)
        ax.set_ylabel('Frequency')
        ax.set_title(f'Histogram of {col}')
        ax.legend(loc='upper right')
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_{col}_hist.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figure_files.append({'path': fname, 'type': 'histogram', 'variable': col, 'num': fig_num})
        plt.close()
        fig_num += 1
        
        # Boxplot
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        bp = ax.boxplot(data, vert=True, patch_artist=True, widths=0.5)
        bp['boxes'][0].set_facecolor(colors['light'])
        bp['boxes'][0].set_edgecolor(colors['primary'])
        bp['medians'][0].set_color(colors['secondary'])
        bp['medians'][0].set_linewidth(2)
        
        # Add statistics text
        stats_text = f'N: {len(data)}\nMean: {data.mean():.2f}\nStd: {data.std():.2f}'
        ax.text(1.3, data.median(), stats_text, fontsize=9, verticalalignment='center',
               bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
        ax.set_ylabel(col)
        ax.set_title(f'Boxplot of {col}')
        ax.set_xticks([])
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_{col}_box.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figure_files.append({'path': fname, 'type': 'boxplot', 'variable': col, 'num': fig_num})
        plt.close()
        fig_num += 1
    
    # 3. Time series (if applicable) - Representative samples
    if var_class['id_vars'] and var_class['time_vars'] and all_numeric:
        id_col = var_class['id_vars'][0]
        time_col = var_class['time_vars'][0]
        subjects = df[id_col].unique()[:config['n_representative']]
        
        for col in all_numeric[:1]:  # One representative
            fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
            for pid in subjects:
                subset = df[df[id_col] == pid]
                ax.plot(subset[time_col], subset[col], label=str(pid), alpha=0.8, lw=1.5)
            ax.set_xlabel('Time')
            ax.set_ylabel(col)
            ax.set_title(f'Time Series of {col} (n={len(subjects)} subjects)')
            ax.legend(loc='upper right')
            ax.grid(True, alpha=0.3)
            plt.tight_layout()
            fname = os.path.join(output_dir, f'fig{fig_num:02d}_{col}_timeseries.png')
            plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
            figure_files.append({'path': fname, 'type': 'timeseries', 'variable': col, 'num': fig_num})
            plt.close()
            fig_num += 1
    
    # 4. Missing value chart
    missing_df = analyze_missing(df)
    if missing_df['n_missing'].sum() > 0:
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        missing_sorted = missing_df.sort_values('pct_missing', ascending=True)
        ax.barh(missing_sorted['variable'], missing_sorted['pct_missing'], color=colors['secondary'])
        ax.set_xlabel('Missing %')
        ax.set_title('Missing Values by Variable')
        ax.set_xlim(0, max(100, missing_sorted['pct_missing'].max() * 1.1))
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_missing_values.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figure_files.append({'path': fname, 'type': 'missing', 'variable': 'all', 'num': fig_num})
        plt.close()
    
    return figure_files


def create_word_document(df, var_class, var_descriptions, missing_df, unique_keys_df, 
                         figure_files, config, output_path):
    """Create comprehensive Word document report."""
    doc = Document()
    
    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading('Exploratory Data Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ================================================================
    # 1. DATASET OVERVIEW
    # ================================================================
    doc.add_heading('1. Dataset Overview', level=1)
    
    doc.add_paragraph(
        f'The dataset contains {len(df):,} observations (rows) and {len(df.columns)} variables (columns). '
        f'Table 1 presents the basic characteristics of the dataset.'
    )
    
    doc.add_paragraph('Table 1. Dataset Characteristics', style='Caption')
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    
    id_var = var_class['id_vars'][0] if var_class['id_vars'] else None
    n_subjects = df[id_var].nunique() if id_var else 'N/A'
    time_var = var_class['time_vars'][0] if var_class['time_vars'] else None
    n_timepoints = df[time_var].nunique() if time_var else 'N/A'
    
    overview_data = [
        ('Total Observations', f'{len(df):,}'),
        ('Number of Variables', str(len(df.columns))),
        ('Number of Subjects', str(n_subjects)),
        ('Number of Time Points', str(n_timepoints))
    ]
    
    for i, (label, value) in enumerate(overview_data):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # ================================================================
    # 2. VARIABLE DESCRIPTIONS
    # ================================================================
    doc.add_heading('2. Variable Descriptions', level=1)
    
    doc.add_paragraph(
        'Table 2 provides a description of each variable in the dataset, including data type and classification.'
    )
    
    doc.add_paragraph('Table 2. Variable Descriptions', style='Caption')
    table2 = doc.add_table(rows=len(df.columns)+1, cols=4)
    table2.style = 'Table Grid'
    
    headers = ['Variable', 'Type', 'Classification', 'Description']
    for j, h in enumerate(headers):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    
    for i, col in enumerate(df.columns):
        # Determine classification
        if col in var_class['id_vars']:
            classification = 'ID'
        elif col in var_class['time_vars']:
            classification = 'Time'
        elif col in var_class['categorical']:
            classification = 'Categorical'
        elif col in var_class['numeric_discrete']:
            classification = 'Numeric (Discrete)'
        elif col in var_class['numeric_continuous']:
            classification = 'Numeric (Continuous)'
        else:
            classification = 'Unknown'
        
        table2.rows[i+1].cells[0].text = col
        table2.rows[i+1].cells[1].text = str(df[col].dtype)
        table2.rows[i+1].cells[2].text = classification
        table2.rows[i+1].cells[3].text = var_descriptions.get(col, '-')
    
    doc.add_paragraph()
    
    # ================================================================
    # 3. MISSING VALUE ANALYSIS
    # ================================================================
    doc.add_heading('3. Missing Value Analysis', level=1)
    
    total_missing = missing_df['n_missing'].sum()
    if total_missing == 0:
        doc.add_paragraph('No missing values were detected in the dataset.')
    else:
        doc.add_paragraph(
            f'A total of {total_missing:,} missing values were identified across all variables. '
            f'Table 3 shows the missing value statistics for each variable.'
        )
        
        doc.add_paragraph('Table 3. Missing Value Analysis', style='Caption')
        table3 = doc.add_table(rows=len(missing_df)+1, cols=4)
        table3.style = 'Table Grid'
        
        headers = ['Variable', 'N Valid', 'N Missing', 'Missing %']
        for j, h in enumerate(headers):
            table3.rows[0].cells[j].text = h
            table3.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        
        for i, row in missing_df.iterrows():
            table3.rows[i+1].cells[0].text = row['variable']
            table3.rows[i+1].cells[1].text = str(row['n_valid'])
            table3.rows[i+1].cells[2].text = str(row['n_missing'])
            table3.rows[i+1].cells[3].text = f"{row['pct_missing']:.2f}%"
        
        # Add missing value figure if exists
        missing_fig = [f for f in figure_files if f['type'] == 'missing']
        if missing_fig:
            doc.add_paragraph()
            doc.add_picture(missing_fig[0]['path'], width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {missing_fig[0]['num']}. Missing values by variable", style='Caption')
    
    doc.add_paragraph()
    
    # ================================================================
    # 4. UNIQUE KEY ANALYSIS
    # ================================================================
    doc.add_heading('4. Unique Key Analysis', level=1)
    
    unique_keys = unique_keys_df[unique_keys_df['is_unique_key'] == True]
    if len(unique_keys) > 0:
        doc.add_paragraph(
            f'The following column(s) uniquely identify each observation: {", ".join(unique_keys["columns"].tolist())}. '
            f'Table 4 shows the uniqueness analysis for all potential keys.'
        )
    else:
        doc.add_paragraph(
            'No single column or combination uniquely identifies each observation. '
            'Table 4 shows the uniqueness analysis.'
        )
    
    doc.add_paragraph('Table 4. Unique Key Analysis', style='Caption')
    table4 = doc.add_table(rows=len(unique_keys_df)+1, cols=4)
    table4.style = 'Table Grid'
    
    headers = ['Column(s)', 'Unique Values', 'Coverage %', 'Is Unique Key']
    for j, h in enumerate(headers):
        table4.rows[0].cells[j].text = h
        table4.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    
    for i, row in unique_keys_df.iterrows():
        table4.rows[i+1].cells[0].text = str(row['columns'])
        table4.rows[i+1].cells[1].text = str(row['n_unique'])
        table4.rows[i+1].cells[2].text = f"{row['coverage']:.1f}%"
        table4.rows[i+1].cells[3].text = 'Yes' if row['is_unique_key'] else 'No'
    
    doc.add_paragraph()
    
    # ================================================================
    # 5. ID/TIME VARIABLE CHARACTERISTICS
    # ================================================================
    if var_class['id_vars'] or var_class['time_vars']:
        doc.add_heading('5. ID and Time Variable Characteristics', level=1)
        
        doc.add_paragraph(
            'ID and Time variables are excluded from descriptive statistics but their format is described below.'
        )
        
        for col in var_class['id_vars']:
            doc.add_paragraph(f"• {col} (ID): {df[col].nunique()} unique values, "
                            f"format: {df[col].dtype}, sample: {df[col].unique()[:3].tolist()}")
        
        for col in var_class['time_vars']:
            doc.add_paragraph(f"• {col} (Time): range {df[col].min()} to {df[col].max()}, "
                            f"{df[col].nunique()} unique values")
        
        doc.add_paragraph()
    
    # ================================================================
    # 6. CATEGORICAL VARIABLE ANALYSIS
    # ================================================================
    table_num = 5
    
    if var_class['categorical']:
        doc.add_heading('6. Categorical Variable Analysis', level=1)
        
        for col in var_class['categorical']:
            doc.add_heading(f'6.{var_class["categorical"].index(col)+1}. {col}', level=2)
            
            freq_df = compute_categorical_stats(df[col])
            
            doc.add_paragraph(
                f'The variable "{col}" has {len(freq_df)} categories. '
                f'The most frequent category is "{freq_df.iloc[0]["Category"]}" '
                f'({freq_df.iloc[0]["Count"]} observations, {freq_df.iloc[0]["Percentage"]:.1f}%). '
                f'Table {table_num} presents the frequency distribution.'
            )
            
            # Frequency Table
            doc.add_paragraph(f'Table {table_num}. Frequency Distribution of {col}', style='Caption')
            tbl = doc.add_table(rows=len(freq_df)+1, cols=3)
            tbl.style = 'Table Grid'
            
            headers = ['Category', 'Count (n)', 'Percentage (%)']
            for j, h in enumerate(headers):
                tbl.rows[0].cells[j].text = h
                tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            
            for idx, row in freq_df.iterrows():
                tbl.rows[idx+1].cells[0].text = str(row['Category'])
                tbl.rows[idx+1].cells[1].text = str(row['Count'])
                tbl.rows[idx+1].cells[2].text = f"{row['Percentage']:.1f}"
            
            table_num += 1
            doc.add_paragraph()
            
            # Figures
            pie_fig = [f for f in figure_files if f['variable'] == col and f['type'] == 'pie']
            bar_fig = [f for f in figure_files if f['variable'] == col and f['type'] == 'bar']
            
            if pie_fig:
                doc.add_picture(pie_fig[0]['path'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Figure {pie_fig[0]['num']}. Pie chart of {col}", style='Caption')
            
            if bar_fig:
                doc.add_picture(bar_fig[0]['path'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Figure {bar_fig[0]['num']}. Frequency distribution of {col}", style='Caption')
            
            doc.add_paragraph()
    
    # ================================================================
    # 7. NUMERIC VARIABLE ANALYSIS
    # ================================================================
    all_numeric = var_class['numeric_discrete'] + var_class['numeric_continuous']
    
    if all_numeric:
        doc.add_heading('7. Numeric Variable Analysis', level=1)
        
        # Summary table for all numeric variables
        doc.add_paragraph(
            f'Table {table_num} presents the descriptive statistics for all numeric variables.'
        )
        
        doc.add_paragraph(f'Table {table_num}. Descriptive Statistics for Numeric Variables', style='Caption')
        
        stats_list = []
        for col in all_numeric:
            stats = compute_numeric_stats(df[col])
            stats['Variable'] = col
            stats_list.append(stats)
        
        stats_df = pd.DataFrame(stats_list)
        cols_order = ['Variable', 'N', 'Missing', 'Missing %', 'Mean', 'Std', 'Min', '25%', '50% (Median)', '75%', 'Max']
        
        tbl = doc.add_table(rows=len(stats_df)+1, cols=len(cols_order))
        tbl.style = 'Table Grid'
        
        for j, h in enumerate(cols_order):
            tbl.rows[0].cells[j].text = h
            tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        
        for i, (idx, row) in enumerate(stats_df.iterrows()):
            for j, col_name in enumerate(cols_order):
                val = row[col_name]
                if isinstance(val, float):
                    tbl.rows[i+1].cells[j].text = f"{val:.2f}"
                else:
                    tbl.rows[i+1].cells[j].text = str(val)
        
        table_num += 1
        doc.add_paragraph()
        
        # Individual variable analysis with figures
        for col in all_numeric:
            var_type = 'Discrete' if col in var_class['numeric_discrete'] else 'Continuous'
            doc.add_heading(f'7.{all_numeric.index(col)+1}. {col} ({var_type})', level=2)
            
            stats = compute_numeric_stats(df[col])
            
            doc.add_paragraph(
                f'The variable "{col}" has {stats["N"]} valid observations with {stats["Missing"]} missing values '
                f'({stats["Missing %"]:.1f}%). The mean value is {stats["Mean"]:.2f} (SD: {stats["Std"]:.2f}), '
                f'with a median of {stats["50% (Median)"]:.2f}. Values range from {stats["Min"]:.2f} to {stats["Max"]:.2f}.'
            )
            
            # Histogram
            hist_fig = [f for f in figure_files if f['variable'] == col and f['type'] == 'histogram']
            if hist_fig:
                doc.add_picture(hist_fig[0]['path'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Figure {hist_fig[0]['num']}. Histogram of {col}", style='Caption')
            
            # Boxplot
            box_fig = [f for f in figure_files if f['variable'] == col and f['type'] == 'boxplot']
            if box_fig:
                doc.add_picture(box_fig[0]['path'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Figure {box_fig[0]['num']}. Boxplot of {col}", style='Caption')
            
            doc.add_paragraph()
    
    # ================================================================
    # 8. TIME SERIES (if applicable)
    # ================================================================
    ts_figs = [f for f in figure_files if f['type'] == 'timeseries']
    if ts_figs:
        doc.add_heading('8. Time Series Visualization', level=1)
        
        doc.add_paragraph(
            f'Figure {ts_figs[0]["num"]} shows the time series pattern for representative subjects.'
        )
        
        doc.add_picture(ts_figs[0]['path'], width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Figure {ts_figs[0]['num']}. Time series of {ts_figs[0]['variable']}", style='Caption')
    
    doc.save(output_path)
    return output_path


def generate_report(df, output='analysis_report.docx', api_key=None, config=None,
                    verbose=True, download=False):
    """
    Generate a publication-ready EDA report from a DataFrame.
    
    Parameters
    ----------
    df : pandas.DataFrame
        Input data
    output : str
        Output filename (default: 'analysis_report.docx')
    api_key : str, optional
        OpenAI API key
    config : dict, optional
        Configuration overrides
    verbose : bool
        Print progress (default: True)
    download : bool
        Auto-download in Colab (default: False)
    
    Returns
    -------
    str
        Path to generated zip file
    """
    
    cfg = DEFAULT_CONFIG.copy()
    if config:
        cfg.update(config)
    
    if verbose:
        print("=" * 50)
        print("EDA REPORT GENERATOR v2.0")
        print("=" * 50)
    
    # 1. OpenAI client
    if verbose:
        print("\n[1/7] Connecting to OpenAI API...")
    client = _get_openai_client(api_key)
    
    # 2. Variable descriptions
    if verbose:
        print("[2/7] Generating variable descriptions...")
    var_descriptions = get_variable_descriptions(df, client)
    
    # 3. Variable classification
    if verbose:
        print("[3/7] Classifying variables...")
    var_class = classify_variables(df, client)
    
    # 4. Missing value analysis
    if verbose:
        print("[4/7] Analyzing missing values...")
    missing_df = analyze_missing(df)
    
    # 5. Unique key analysis
    if verbose:
        print("[5/7] Analyzing unique keys...")
    unique_keys_df = analyze_unique_keys(df, var_class)
    
    # 6. Generate figures
    if verbose:
        print("[6/7] Generating figures...")
    output_dir = os.path.dirname(output) or '.'
    figure_files = generate_figures(df, var_class, cfg, output_dir)
    
    # 7. Create Word document
    if verbose:
        print("[7/7] Creating Word document...")
    doc_path = create_word_document(df, var_class, var_descriptions, missing_df,
                                     unique_keys_df, figure_files, cfg, output)
    
    # Create zip file
    zip_filename = output.replace('.docx', '.zip')
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipf.write(doc_path, os.path.basename(doc_path))
        for f in figure_files:
            if os.path.exists(f['path']):
                zipf.write(f['path'], os.path.basename(f['path']))
    
    # Clean up
    if os.path.exists(doc_path):
        os.remove(doc_path)
    for f in figure_files:
        if os.path.exists(f['path']):
            os.remove(f['path'])
    
    if verbose:
        print("\n" + "=" * 50)
        print(f"COMPLETE: {zip_filename}")
        print("=" * 50)
    
    if download:
        try:
            from google.colab import files
            files.download(zip_filename)
        except:
            pass
    
    return zip_filename
